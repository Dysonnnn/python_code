#用python-docx修改已存在的Word文档的表格的字体格式：
#参考链接 https://blog.csdn.net/a15986714591/article/details/78212187?locationNum=4&fps=1

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
document = Document('word.docx')
tables = document.tables[0]

run = tables.cell(1,0).paragraphs[0].add_wordrun('smida')
run.font.name = '黑体'
#run.font.size = 140000
#tables.cell(1,0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

document.save('word_output.docx')
